# app.py
import os
import json
from datetime import datetime, timedelta
from io import BytesIO
from functools import wraps

from flask import Flask, request, jsonify, send_file, render_template, g
from flask_cors import CORS
from sqlalchemy import text
from werkzeug.utils import secure_filename
from threading import Thread
from time import sleep

import bcrypt
import jwt

from openpyxl import Workbook
from docx import Document as DocxDocument
import smtplib
from email.message import EmailMessage
from docx.shared import Inches
from flask import send_file

from models import db, UploadedFile, Document, SelectedEntry, SearchHistory, User
from extractor import extract_rows_from_excel

# -----------------------------
# Render Deployment Paths
# -----------------------------
RENDER = os.environ.get("RENDER", "false") == "true"

if RENDER:
    DB_ROOT = "/opt/render/project/data"
else:
    DB_ROOT = os.path.abspath(".")

UPLOAD_FOLDER = os.path.join(DB_ROOT, "uploads")
DB_PATH = os.path.join(DB_ROOT, "Adoodle.db")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXT = {'xls', 'xlsx', 'csv'}

JWT_SECRET = os.environ.get('JWT_SECRET', 'change-me-please')
JWT_ALGORITHM = os.environ.get('JWT_ALGORITHM', 'HS256')
JWT_EXP_DAYS = int(os.environ.get('JWT_EXP_DAYS', '7'))

SUPER_ADMIN_EMAIL = os.environ.get('SUPER_ADMIN_EMAIL', 'admin@example.com')
SUPER_ADMIN_PW = os.environ.get('SUPER_ADMIN_PW', 'admin123')

def create_app(db_path=f"sqlite:///{DB_PATH}"):
    app = Flask(__name__, static_folder='static')
    app.config.update(
        SQLALCHEMY_DATABASE_URI=db_path,
        SQLALCHEMY_TRACK_MODIFICATIONS=False,
        UPLOAD_FOLDER=UPLOAD_FOLDER
    )

    db.init_app(app)

    with app.app_context():
        db.create_all()
        # create FTS virtual table if possible
        try:
            db.session.execute(text("""
                CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts 
                USING fts5(
                    docid UNINDEXED, 
                    purchasername, 
                    sellername, 
                    propertydescription, 
                    docname, 
                    docno, 
                    content=''
                );
            """))
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(" Warning: SQLite FTS5 not available:", e)

        # Ensure a super-admin exists
        try:
            existing_admin = User.query.filter_by(is_admin=True).first()
            if not existing_admin:
                pw = SUPER_ADMIN_PW
                hashed = bcrypt.hashpw(pw.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                # leave expiry_date None for super-admin
                admin = User(email=SUPER_ADMIN_EMAIL.lower(), password_hash=hashed, name="Super Admin", is_admin=True, is_active=True)
                db.session.add(admin)
                db.session.commit()
                print(f" Created super admin: {SUPER_ADMIN_EMAIL} (change password immediately)")
        except Exception as e:
            db.session.rollback()
            print(" Could not create super admin:", e)

    return app

def cleanup_old_files():
    from app import app  # important: get the running app instance

    with app.app_context():  # <<< FIXED
        try:
            cutoff = datetime.utcnow() - timedelta(days=30)

            old_files = UploadedFile.query.filter(
                UploadedFile.upload_date < cutoff
            ).all() 

            if not old_files:
                return

            for f in old_files:

                # delete file from disk
                try:
                    if os.path.exists(f.filepath):
                        os.remove(f.filepath)
                except Exception as e:
                    print("File delete error:", e)

                # delete DB entry
                db.session.delete(f)

            db.session.commit()
            print(f"Cleanup: {len(old_files)} old files removed.")

        except Exception as e:
            print("Cleanup error:", str(e))

def start_cleanup_thread():
    def worker():
        while True:
            cleanup_old_files()
            sleep(24 * 60 * 60)   # run once per day
    Thread(target=worker, daemon=True).start()

app = create_app()
CORS(app, resources={
    r"/*": {
        "origins": [
            "https://adoodle-software.onrender.com",
            "http://localhost:5173"
        ],
        "supports_credentials": True,
        "allow_headers": ["Authorization", "Content-Type"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"]
    }
})


# ---------------- HELPERS ----------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXT


def build_in_params(int_ids, prefix="id"):
    placeholders = ", ".join([f":{prefix}{i}" for i in range(len(int_ids))])
    params = {f"{prefix}{i}": int_ids[i] for i in range(len(int_ids))}
    return placeholders, params


# ---------------- JWT HELPERS & DECORATOR ----------------
def create_token(user, expires_days=None):
    if expires_days is None:
        expires_days = JWT_EXP_DAYS
    exp = datetime.utcnow() + timedelta(days=expires_days)
    payload = {
        'sub': str(user.id),
        'email': user.email,
        'name': user.name,
        'role': 'admin' if user.is_admin else 'user',
        'exp': exp,
        'iat': datetime.utcnow()
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return token if isinstance(token, str) else token.decode()


def decode_token(token):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except Exception:
        return None


def get_token_from_header():
    auth = request.headers.get('Authorization', '')
    if auth.startswith('Bearer '):
        return auth.split(' ', 1)[1].strip()
    if request.is_json:
        try:
            return request.get_json().get('token')
        except Exception:
            return None
    return request.args.get('token')


def _check_and_handle_expiry(user):
    """
    If user has expiry_date and it's passed, mark as inactive and return True (expired).
    Returns True if expired (and user.is_active updated), False otherwise.
    """
    try:
        expiry = getattr(user, 'expiry_date', None)
    except Exception:
        expiry = None

    if expiry and isinstance(expiry, datetime):
        if datetime.utcnow() > expiry:
            if user.is_active:
                user.is_active = False
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            return True
    return False


def jwt_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        token = get_token_from_header()
        if not token:
            return jsonify({'error': 'Authorization token required'}), 401
        payload = decode_token(token)
        if not payload:
            return jsonify({'error': 'Invalid or expired token'}), 401
        user_id = payload.get('sub')
        try:
            user = User.query.get(int(user_id))
        except Exception:
            user = None
        if not user:
            return jsonify({'error': 'User not found'}), 401

        # Check expiry: if expired, auto-deactivate and deny access
        if _check_and_handle_expiry(user):
            return jsonify({'error': 'Account expired. Contact admin.'}), 403

        if not user.is_active:
            return jsonify({'error': 'User inactive. Contact admin.'}), 403

        g.current_user = user
        return f(*args, **kwargs)
    return wrapper


def admin_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        token = get_token_from_header()
        if not token:
            return jsonify({'error': 'Authorization token required'}), 401
        payload = decode_token(token)
        if not payload:
            return jsonify({'error': 'Invalid or expired token'}), 401
        if payload.get('role') != 'admin':
            return jsonify({'error': 'Admin access required'}), 403
        user_id = payload.get('sub')
        try:
            user = User.query.get(int(user_id))
        except Exception:
            user = None
        if not user or not user.is_admin:
            return jsonify({'error': 'Admin not found'}), 403

        # Check expiry for admin too (optional; can remove if admins shouldn't expire)
        if _check_and_handle_expiry(user):
            return jsonify({'error': 'Admin account expired'}), 403

        g.current_user = user
        return f(*args, **kwargs)
    return wrapper


@app.route('/login', methods=['POST'])
def login():
    data = request.get_json(force=True)
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'Invalid credentials'}), 401

    # If user has expiry_date and it's passed, auto-deactivate and deny login
    if _check_and_handle_expiry(user):
        return jsonify({'error': 'User account expired. Contact admin.'}), 403

    # check active status
    if not user.is_active:
        return jsonify({'error': 'User inactive. Contact admin.'}), 403

    try:
        ok = bcrypt.checkpw(password.encode('utf-8'), user.password_hash.encode('utf-8'))
    except Exception:
        ok = False

    if not ok:
        return jsonify({'error': 'Invalid credentials'}), 401

    token = create_token(user)
    return jsonify({'message': 'ok', 'token': token, 'user': user.as_dict()})


@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    data = request.get_json(force=True)
    email = (data.get('email') or '').strip().lower()
    if not email:
        return jsonify({'error': 'Email required'}), 400

    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'message': 'If an account exists, reset instructions will be sent'}), 200

    payload = {'sub': str(user.id), 'iat': datetime.utcnow(), 'exp': datetime.utcnow() + timedelta(hours=1)}
    token = jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)
    token = token if isinstance(token, str) else token.decode()
    return jsonify({'message': 'reset token generated (dev)', 'reset_token': token}), 200

@app.route('/reset-password', methods=['POST'])
def reset_password():
    data = request.get_json(force=True)
    token = data.get('token')
    new_password = data.get('new_password')
    if not token or not new_password:
        return jsonify({'error': 'Token and new_password required'}), 400

    payload = decode_token(token)
    if not payload:
        return jsonify({'error': 'Invalid or expired token'}), 400

    user_id = payload.get('sub')
    user = User.query.get(int(user_id)) if user_id else None
    if not user:
        return jsonify({'error': 'User not found'}), 404

    # Hash the new password
    hashed = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    user.password_hash = hashed
    db.session.commit()

    return jsonify({'message': 'Password reset successful.'}), 200



# ---------------- Admin APIs ----------------
@app.route('/admin/users', methods=['GET'])
@admin_required
def admin_list_users():
    rows = User.query.order_by(User.created_at.desc()).all()
    out = [u.as_dict() for u in rows]
    return jsonify({'users': out}), 200


@app.route('/admin/create_user', methods=['POST'])
@admin_required
def admin_create_user():
    data = request.get_json(force=True)
    email = (data.get('email') or '').strip().lower()
    password = data.get('password')
    name = data.get('name', None)
    is_admin = bool(data.get('is_admin', False))
    is_active = True if data.get('is_active') is None else bool(data.get('is_active'))
    phone1 = data.get("phone1", None)
    phone2 = data.get("phone2", None)
    address = data.get("address", None)
    
    # optional expiry_date for new user
    expiry_date_str = data.get('expiry_date')
    expiry_date = None
    if expiry_date_str:
        try:
            expiry_date = datetime.fromisoformat(expiry_date_str)
        except Exception:
            return jsonify({'error': 'Invalid expiry_date format. Use ISO format like 2025-12-31T23:59:59'}), 400

    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    if User.query.filter_by(email=email).first():
        return jsonify({'error': 'Email already registered'}), 400

    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    user = User(email=email, password_hash=hashed, name=name, is_admin=is_admin, is_active=is_active,phone1=phone1,
    phone2=phone2,
    address=address)
    if hasattr(user, 'expiry_date'):
        user.expiry_date = expiry_date

    db.session.add(user)
    db.session.commit()
    return jsonify({'message': 'created', 'user': user.as_dict()}), 201


@app.route('/admin/set_status', methods=['POST'])
@admin_required
def admin_set_status():
    data = request.get_json(force=True)
    user_id = data.get('user_id')
    status = data.get('status')
    try:
        user_id = int(user_id)
    except Exception:
        return jsonify({'error': 'Invalid user_id'}), 400
    u = User.query.get(user_id)
    if not u:
        return jsonify({'error': 'User not found'}), 404
    # prevent admin disabling themselves
    admin_user = g.current_user
    if u.id == admin_user.id:
        return jsonify({'error': 'Cannot change your own status'}), 400
    u.is_active = bool(status)
    db.session.commit()
    return jsonify({'message': 'ok', 'user': u.as_dict()}), 200


# ---------------- NEW: Admin set expiry (recharge / validity) ----------------
@app.route('/admin/set_expiry', methods=['POST'])
@admin_required
def admin_set_expiry():
    """
    Body: { user_id: <int>, expiry_date: <ISO datetime string or null> }
    expiry_date example: "2025-12-31T23:59:59" or null to remove expiry
    """
    data = request.get_json(force=True)
    user_id = data.get('user_id')
    expiry_date_str = data.get('expiry_date', None)

    try:
        user_id = int(user_id)
    except Exception:
        return jsonify({'error': 'Invalid user_id'}), 400

    u = User.query.get(user_id)
    if not u:
        return jsonify({'error': 'User not found'}), 404

    # prevent admin changing their own expiry to avoid locking themselves out (optional)
    if u.id == g.current_user.id:
        return jsonify({'error': 'Cannot change your own expiry'}), 400

    if expiry_date_str is None or expiry_date_str == "":
        # Clear expiry (no limit)
        if hasattr(u, 'expiry_date'):
            u.expiry_date = None
            # optionally keep is_active unchanged, or set True
            # u.is_active = True
    else:
        try:
            new_expiry = datetime.fromisoformat(expiry_date_str)
        except Exception:
            return jsonify({'error': 'Invalid expiry_date format, use ISO format'}), 400

        if hasattr(u, 'expiry_date'):
            u.expiry_date = new_expiry
        # If expiry is in past, mark inactive immediately
        if new_expiry and datetime.utcnow() > new_expiry:
            u.is_active = False
        else:
            # if admin sets a future expiry, ensure user is active (option)
            u.is_active = True

    db.session.commit()
    return jsonify({'message': 'expiry updated', 'user': u.as_dict()}), 200

@app.route("/admin/update_user", methods=["POST"])
@admin_required
def admin_update_user():
    data = request.get_json()

    user_id = data.get("id")
    user = User.query.get(user_id)

    if not user:
        return jsonify({"error": "User not found"}), 404

    # --- BASIC FIELDS ---
    user.name = data.get("name", user.name)
    user.email = data.get("email", user.email)
    user.phone1 = data.get("phone1", user.phone1)
    user.phone2 = data.get("phone2", user.phone2)
    user.address = data.get("address", user.address)

    # --- ACTIVE STATUS ---
    if "is_active" in data:
        user.is_active = bool(data.get("is_active"))

    # --- ADMIN ROLE ---
    if "is_admin" in data:
        user.is_admin = bool(data.get("is_admin"))
    
    # --- EXPIRY DATE ---
    expiry_date_str = data.get("expiry_date")
    if expiry_date_str:
        try:
            user.expiry_date = datetime.fromisoformat(expiry_date_str)
        except Exception:
            return jsonify({"error": "Invalid expiry_date format"}), 400

    # --- PASSWORD RESET (ADMIN ONLY) ---
    new_password = data.get("new_password")
    if new_password and new_password.strip() != "":
        hashed = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode('utf-8')
        user.password_hash = hashed

    db.session.commit()

    return jsonify({
        "message": "User updated successfully",
        "user": {
            "id": user.id,
            "email": user.email,
            "name": user.name,
            "phone1": user.phone1,
            "phone2": user.phone2,
            "address": user.address,
            "created_at": user.created_at.isoformat(),
            "is_admin": user.is_admin,
            "is_active": user.is_active,
            "expiry_date": user.expiry_date.isoformat() if user.expiry_date else None
        }
    }), 200

@app.route('/admin/delete_user', methods=['POST'])
@admin_required
def admin_delete_user():
    data = request.get_json()
    user_id = data.get("user_id")

    # Get current logged-in admin ID (adjust if using Flask-Login / JWT)
    current_admin_id = g.current_user.id  

    # Prevent deleting own account
    if int(user_id) == int(current_admin_id):
        return jsonify({"error": "You cannot delete your own account."}), 400

    user = User.query.get(user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    db.session.delete(user)
    db.session.commit()

    return jsonify({"message": "User deleted successfully"}), 200



# ---------------- PROFILE (protected) ----------------
@app.route('/profile', methods=['GET'])
@jwt_required
def get_profile():
    user = g.current_user
    return jsonify({
        'id': user.id,
        'email': user.email,
        'name': user.name,
        'is_admin': user.is_admin,
        'is_active': user.is_active,
        'created_at': user.created_at.isoformat(),
        'expiry_date': getattr(user, 'expiry_date', None).isoformat() if getattr(user, 'expiry_date', None) else None
    }), 200


@app.route('/profile/update', methods=['PUT'])
@jwt_required
def update_profile():
    user = g.current_user
    data = request.json

    new_name = data.get("name")
    new_email = data.get("email")
    old_password = data.get("old_password")
    new_password = data.get("new_password")

    updated = False

    if new_name and new_name.strip() != "" and new_name != user.name:
        user.name = new_name
        updated = True

    if new_email and new_email != user.email:
        exists = User.query.filter_by(email=new_email).first()
        if exists:
            return jsonify({"error": "Email already exists"}), 400
        user.email = new_email
        updated = True

    if new_password:
        if not old_password:
            return jsonify({"error": "Old password is required"}), 400

        try:
            ok = bcrypt.checkpw(old_password.encode(), user.password_hash.encode('utf-8'))
        except Exception:
            ok = False
        if not ok:
            return jsonify({"error": "Old password is incorrect"}), 400

        hashed = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode('utf-8')
        user.password_hash = hashed
        updated = True

    if not updated:
        return jsonify({"message": "No changes made"}), 200

    db.session.commit()
    return jsonify({"message": "Profile updated successfully"}), 200

# ---------------- API: fetch selected rows (by ids) ----------------
@app.route('/api/selected_rows', methods=['POST'])
@jwt_required
def api_selected_rows():
    payload = request.get_json(force=True)
    ids = payload.get('ids', []) or []

    user_id = g.current_user.id

    # Fetch selected entries (only this user's)
    if not ids:
        sel_rows = SelectedEntry.query.filter_by(user_id=user_id).order_by(SelectedEntry.created_at.desc()).all()
    else:
        try:
            int_ids = [int(i) for i in ids]
        except:
            return jsonify({'error': 'Invalid ids'}), 400

        sel_rows = SelectedEntry.query.filter(
            SelectedEntry.user_id == user_id,
            SelectedEntry.document_id.in_(int_ids)
        ).all()

    if not sel_rows:
        return jsonify({'groups': []})

    doc_ids = [s.document_id for s in sel_rows]
    sel_map = {s.document_id: s for s in sel_rows}

    # Fetch document rows that belong to this user
    placeholders, params = build_in_params(doc_ids, prefix="id")
    params['user_id'] = user_id
    sql = text(f"""
        SELECT id, table_name, docno, docname, registrationdate, sroname,
               sellername, purchasername, propertydescription, areaname,
               consideration_amt, dateofexecution
        FROM documents
        WHERE id IN ({placeholders}) AND user_id = :user_id
        ORDER BY id DESC
    """)
    rows = db.session.execute(sql, params).fetchall()

    # GROUP BY table_name
    groups = {}
    for r in rows:
        doc_id = r[0]
        table_name = r[1]

        if table_name not in groups:
            groups[table_name] = {
                "table_name": table_name,
                "chip_label": table_name,
                "rows": []
            }

        groups[table_name]["rows"].append({
            "sel_id": sel_map.get(doc_id).id,
            "document_id": doc_id,
            "docno": r[2],
            "docname": r[3],
            "registrationdate": r[4],
            "sroname": r[5],
            "sellerparty": r[6],
            "purchaserparty": r[7],
            "propertydescription": r[8],
            "areaname": r[9],
            "consideration_amt": r[10],
            "dateofexecution": r[11]
        })

    return jsonify({
        "groups": list(groups.values())
    })


# ---------------- API: save selected entries (persist) ----------------
@app.route('/api/save_selected', methods=['POST'])
@jwt_required
def api_save_selected():
    payload = request.get_json(force=True)
    entries = payload.get('entries', [])  # expect list of {id}
    user_id = g.current_user.id

    if not entries:
        return jsonify({'error': 'No entries provided'}), 400

    doc_ids = []
    for e in entries:
        try:
            doc_id = int(e.get('id'))
        except:
            continue
        doc_ids.append(doc_id)

    if not doc_ids:
        return jsonify({'error': 'No valid ids'}), 400

    added = 0
    for doc_id in doc_ids:
        # Make sure this document belongs to the current user
        doc = Document.query.filter_by(id=doc_id, user_id=user_id).first()
        if not doc:
            continue

        # Already selected?
        exists = SelectedEntry.query.filter_by(document_id=doc_id, user_id=user_id).first()
        if exists:
            continue

        table_name = doc.table_name

        se = SelectedEntry(
            user_id=user_id,
            document_id=doc_id,
            table_name=table_name,
            label=table_name
        )

        db.session.add(se)
        added += 1

    db.session.commit()

    return jsonify({
        'added': added,
        'saved_total': len(doc_ids)
    })


# ---------------- API: remove selected entry (persisted) ----------------
@app.route('/api/remove_selected', methods=['POST', 'DELETE'])
@jwt_required
def api_remove_selected():
    if request.method == 'DELETE':
        sid = request.args.get('id')
    else:
        payload = request.get_json(force=True)
        sid = payload.get('id')
    try:
        sid = int(sid)
    except Exception:
        return jsonify({'error': 'Invalid id'}), 400

    se = SelectedEntry.query.get(sid)
    if not se or se.user_id != g.current_user.id:
        return jsonify({'error': 'Not found or not yours'}), 404

    db.session.delete(se)
    db.session.commit()
    return jsonify({'deleted': sid})

@app.route('/api/remove_selected_group', methods=['POST'])
@jwt_required
def api_remove_selected_group():
    payload = request.get_json(force=True)
    table_name = payload.get('table_name')
    user_id = g.current_user.id

    if not table_name:
        return jsonify({'error': 'No table name given'}), 400

    # find docs of this user and table
    docs = Document.query.filter_by(table_name=table_name, user_id=user_id).all()
    doc_ids = [d.id for d in docs]

    if not doc_ids:
        return jsonify({'deleted': 0})

    deleted_count = SelectedEntry.query.filter(
        SelectedEntry.user_id == user_id,
        SelectedEntry.document_id.in_(doc_ids)
    ).delete(synchronize_session=False)

    db.session.commit()

    return jsonify({'deleted': deleted_count})


@app.route('/upload', methods=['POST'])
@jwt_required
def upload_files():
    files = request.files.getlist('files')

    table_name = request.form.get("table_name", "").strip()
    if not table_name:
        return jsonify({"status": "error", "message": "Table name is required"}), 400

    # folder per table per user (avoid collisions)
    user_folder = os.path.join(app.config['UPLOAD_FOLDER'], str(g.current_user.id))
    table_folder = os.path.join(user_folder, table_name)
    os.makedirs(table_folder, exist_ok=True)

    if not files:
        return jsonify({'status': 'error', 'message': 'No files uploaded'}), 400

    created_files = []

    def safe_float(v):
        try:
            if v is None or v == "" or str(v).strip() == "":
                return None
            return float(str(v).replace(",", "").strip())
        except:
            return None

    for file in files:
        if not allowed_file(file.filename):
            continue

        try:
            original_name = secure_filename(file.filename)
            fname = original_name
            fpath = os.path.join(table_folder, fname)

            # make unique name
            base, ext = os.path.splitext(fpath)
            i = 1
            while os.path.exists(fpath):
                fname = f"{os.path.splitext(original_name)[0]}_{i}{ext}"
                fpath = os.path.join(table_folder, fname)
                i += 1

            file.save(fpath)
            filesize = os.path.getsize(fpath)

            # save Upload record (include user_id)
            uf = UploadedFile(
                user_id=g.current_user.id,
                filename=fname,
                filepath=fpath,
                filesize=filesize,
                table_name=table_name
            )
            db.session.add(uf)
            db.session.flush()

            rows = extract_rows_from_excel(fpath)
            print(f"Parsed {len(rows)} rows from {fname}")

            for r in rows:
                doc = Document(
                    user_id=g.current_user.id,
                    file_id=uf.id,
                    table_name=table_name,  # ✅ FIX ADDED

                    sr_code=str(r.get('sr_code')).strip() if r.get('sr_code') else None,
                    internal_document_number=str(r.get('internal_document_number')).strip() if r.get('internal_document_number') else None,

                    docno=str(r.get('docno')).strip() if r.get('docno') else None,
                    docname=str(r.get('docname')).strip() if r.get('docname') else None,

                    registrationdate=r.get('registrationdate'),
                    dateofexecution=r.get('dateofexecution'),

                    purchasername=str(r.get('purchasername')).strip() if r.get('purchasername') else None,
                    sellername=str(r.get('sellername')).strip() if r.get('sellername') else None,

                    propertydescription=str(r.get('propertydescription')).strip() if r.get('propertydescription') else None,
                    areaname=str(r.get('areaname')).strip() if r.get('areaname') else None,
                    sroname=str(r.get('sroname')).strip() if r.get('sroname') else None,

                    consideration_amt=safe_float(r.get('consideration_amt')),
                    marketvalue=safe_float(r.get('marketvalue')),

                    raw_json=r.get('raw_json')
                )

                db.session.add(doc)
                db.session.flush()

                # FTS insert (docid numeric stored as string previously)
                try:
                    db.session.execute(text("""
                        INSERT INTO documents_fts(
                            docid, purchasername, sellername, propertydescription, docname, docno
                        ) VALUES (:docid, :purchasername, :sellername, :propertydescription, :docname, :docno)
                    """), {
                        'docid': str(doc.id),
                        'purchasername': doc.purchasername or '',
                        'sellername': doc.sellername or '',
                        'propertydescription': doc.propertydescription or '',
                        'docname': doc.docname or '',
                        'docno': doc.docno or ''
                    })
                except Exception as e:
                    # FTS may not be available or may fail for this DB; skip silently (logged)
                    print("FTS skipped:", e)

            db.session.commit()

            created_files.append({
                'file_id': uf.id,
                'filename': uf.filename,
                'table_name': table_name
            })

        except Exception as e:
            db.session.rollback()
            return jsonify({'status': 'error', 'message': str(e)}), 500

    return jsonify({'status': 'success', 'uploaded': created_files}), 201

## ---------------- SEARCH (protected) ----------------
@app.route('/search', methods=['GET'])
@jwt_required
def search():
    q = request.args.get('q', '').strip()
    purchaser = request.args.get('purchaser', '').strip()
    seller = request.args.get('seller', '').strip()
    docname = request.args.get('docname', '').strip()
    docno = request.args.get('docno', '').strip()
    propdesc = request.args.get('propertydescription', '').strip()
    reg_date = request.args.get('registrationdate', '').strip()
    exact = request.args.get('exact', '0') == '1'
    table_name = request.args.get("table_name", "").strip()

    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 300))
    offset = (page - 1) * per_page

    user_id = g.current_user.id

    base_query = """
        SELECT d.id, d.docno, d.docname, d.registrationdate, d.sroname,
               d.sellername, d.purchasername, d.propertydescription,
               d.areaname, d.consideration_amt
        FROM documents d
    """

    where_clauses, params = ["d.user_id = :user_id"], {"user_id": user_id}

    # FILTER BY TABLE NAME FIRST
    if table_name:
        where_clauses.append("d.table_name = :table_name")
        params["table_name"] = table_name

    if q:
        like_q = f"%{q}%"
        where_clauses.append("""
            (d.purchasername LIKE :like_q OR d.sellername LIKE :like_q OR
             d.propertydescription LIKE :like_q OR d.docname LIKE :like_q OR
             d.docno LIKE :like_q OR d.sroname LIKE :like_q OR d.areaname LIKE :like_q)
        """)
        params['like_q'] = like_q

    def add_filter(field, value, param):
        if not value:
            return
        if exact:
            where_clauses.append(f"d.{field} = :{param}")
            params[param] = value
        else:
            where_clauses.append(f"d.{field} LIKE :{param}")
            params[param] = f"%{value}%"

    add_filter('purchasername', purchaser, 'purchaser')
    add_filter('sellername', seller, 'seller')
    add_filter('docname', docname, 'docname_param')
    add_filter('docno', docno, 'docno_param')
    add_filter('propertydescription', propdesc, 'prop_param')

    if reg_date:
        if exact:
            where_clauses.append("d.registrationdate = :reg_date")
            params['reg_date'] = reg_date
        else:
            where_clauses.append("d.registrationdate LIKE :reg_date_like")
            params['reg_date_like'] = f"%{reg_date}%"

    final_where = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
    params.update({'limit': per_page, 'offset': offset})

    total_stmt = text(f"SELECT COUNT(*) FROM documents d {final_where}")
    total = db.session.execute(total_stmt, params).scalar()

    data_stmt = text(
        base_query + final_where +
        " ORDER BY d.id DESC LIMIT :limit OFFSET :offset"
    )
    rows = db.session.execute(data_stmt, params).fetchall()

    results = [{
        'id': r[0],
        'docno': r[1],
        'docname': r[2],
        'registrationdate': r[3],
        'sroname': r[4],
        'sellerparty': r[5],
        'purchaserparty': r[6],
        'propertydescription': r[7],
        'areaname': r[8],
        'consideration_amt': r[9]
    } for r in rows]

    return jsonify({
        'results': results,
        'total': total,
        'page': page,
        'per_page': per_page
    })

@app.route('/tables', methods=['GET'])
@jwt_required
def list_tables():
    # Only list tables uploaded by this user
    sql = text("""
        SELECT DISTINCT table_name 
        FROM uploaded_files
        WHERE user_id = :user_id
        ORDER BY table_name;
    """)
    rows = db.session.execute(sql, {'user_id': g.current_user.id}).fetchall()
    tables = [r[0] for r in rows]

    return jsonify({"tables": tables})

# ---------------- EXPORT & EMAIL (protected) ----------------
@app.route('/export/selected/excel', methods=['POST'])
@jwt_required
def export_selected_excel():
    data = request.json
    ids = [e['id'] for e in data.get("entries", [])]

    if not ids:
        return jsonify({'error': 'No entries selected'}), 400

    # Ensure ownership
    docs = Document.query.filter(Document.id.in_(ids), Document.user_id == g.current_user.id).all()

    wb = Workbook()
    ws = wb.active
    ws.append(["ID", "Doc No", "Doc Name", "Purchaser", "Seller", "Reg Date", "Area", "Consideration"])

    for d in docs:
        ws.append([
            d.id, d.docno, d.docname, d.purchasername, d.sellername,
            d.registrationdate, d.areaname, d.consideration_amt
        ])

    stream = BytesIO()
    wb.save(stream)
    docx = Document()

    for d in docs:

        # --- English doc name ---
        eng_docname = DOCNAME_MAP.get(d.docname.strip(), d.docname)

        # --- Extract year ---
        year = ""
        if d.registrationdate:
            try:
                year = d.registrationdate.split("-")[0]
            except:
                year = ""

        # --- Extract SRO code (Marathi → English Short code) ---
        # Example: "हवेली 23" → "HVL 23"
        sro_eng = ""
        if d.sroname:
            if "हवेली" in d.sroname:
                num = ''.join([c for c in d.sroname if c.isdigit()])
                sro_eng = f"HVL {num}"

        # --- Title Format ---
        title = f"{year} – {eng_docname} dated {d.registrationdate} (Reg. No {sro_eng} {d.docno}/{year})"

        # Add title
        p = docx.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(13)

        # --- 3 column table ---
        table = docx.add_table(rows=1, cols=3)
        table.autofit = True
    stream.seek(0)

    return send_file(
        stream,
        download_name="selected_entries.xlsx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@app.route('/export/selected/word', methods=['POST'])
@jwt_required
def export_selected_word():
    from docx import Document as WordDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from io import BytesIO

    data = request.json
    ids = [e['id'] for e in data.get("entries", [])]

    docs = Document.query.filter(
        Document.id.in_(ids),
        Document.user_id == g.current_user.id
    ).all()

    if not docs:
        return jsonify({"error": "No matching documents"}), 400

    DOCNAME_MAP = {
        "कंफर्मेशन डीड": "Conformation Deed",
        "करारनामा": " Agreement",
        "गहाणखत": "Mortgage Deed",
        "अ‍ॅफिडेव्हिट": "Affidavit",
        "पॉवर ऑफ अटर्नी": "Power of Attorney",
        "डेव्हलपमेंट अ‍ॅग्रीमेंट": "Development Agreement",
        "असाइनमेंट डीड": "Assignment Deed",
        "कमन्समेंट सर्टिफिकेट": "Commencement Certificate",
        "कम्प्लिशन सर्टिफिकेट": "Completion Certificate",
        "रिलीज डीड": "Release Deed",
        "सोसायटी रजिस्ट्रेशन": "Society Registration",
        "ऑक्युपन्सी सर्टिफिकेट": "Part Occupancy Certificate",
        "नियमितीकरण प्रमाणपत्र": "Regularization Certificate",
        "ना आदेश": "NA Order"
    }

    docx = WordDoc()

    for d in docs:
        eng_docname = DOCNAME_MAP.get(d.docname.strip(), d.docname)

        year = ""
        if d.registrationdate:
            try:
                year = d.registrationdate.split("-")[0]
            except:
                year = ""

        sro_eng = ""
        if d.sroname:
            if "हवेली" in d.sroname:
                num = ''.join([c for c in d.sroname if c.isdigit()])
                sro_eng = f"HVL {num}"

        title = f"{year} – {eng_docname} dated {d.registrationdate} (Reg. No {sro_eng}/ {d.docno}/{year})"

        p = docx.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(13)

        table = docx.add_table(rows=1, cols=3)
        table.autofit = True

        hdr = table.rows[0].cells
        hdr[0].text = d.sellername or ""
        hdr[1].text = d.purchasername or ""
        hdr[2].text = d.propertydescription or ""

        docx.add_paragraph("\n")

    buffer = BytesIO()
    docx.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        download_name="selected_formatted.docx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ---------------- RUN ----------------
if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)    
    start_cleanup_thread() 
    app.run(debug=True)